import streamlit as st
from io import BytesIO
from PIL import Image
import pdfkit
from PyPDF2 import PdfMerger
import requests
import tempfile
import os
import ocrmypdf

BACKEND_URL = "http://backend:8000/process/"

def download_docx(questions):
    import docx
    doc = docx.Document()
    for section, items in questions.items():
        doc.add_heading(section.replace('_', ' ').title(), level=2)
        for item in items:
            if isinstance(item, dict):
                doc.add_paragraph(f"Q: {item.get('question', '')}")
                for opt in item.get("options", []):
                    doc.add_paragraph(f" - {opt}")
            else:
                doc.add_paragraph(f"Q: {item}")
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def download_txt(questions):
    text = ""
    for section, items in questions.items():
        text += f"\n== {section.replace('_', ' ').title()} ==\n"
        for item in items:
            if isinstance(item, dict):
                text += f"Q: {item.get('question', '')}\n"
                for opt in item.get("options", []):
                    text += f" - {opt}\n"
            else:
                text += f"Q: {item}\n"
    return text

def convert_txt_to_pdf(text):
    return pdfkit.from_string(text, False)

def get_ordered_files(uploaded_files):
    def extract_order(file_name):
        base = os.path.splitext(file_name)[0]
        parts = base.split('_')
        try:
            return int(parts[-1])
        except:
            return 9999  # Files without order suffix go at the end
    return sorted(uploaded_files, key=lambda f: extract_order(f.name))

def ocr_pdf_to_searchable_pdf(pdf_bytes, languages="tam+hin+eng"):
    with tempfile.NamedTemporaryFile(suffix=".pdf") as input_tmp, tempfile.NamedTemporaryFile(suffix=".pdf") as output_tmp:
        input_tmp.write(pdf_bytes)
        input_tmp.flush()
        # Use skip_text=True and output_type='pdf' to avoid Ghostscript issues
        ocrmypdf.ocr(
            input_tmp.name,
            output_tmp.name,
            language=languages,
            deskew=True,
            clean=True,
            optimize=3,
            progress_bar=False,
            use_threads=True,
            skip_text=True,
            output_type='pdf'
        )
        with open(output_tmp.name, 'rb') as f:
            searchable_pdf = f.read()
    return BytesIO(searchable_pdf)

st.set_page_config(page_title="📘 Question Bank Generator", layout="wide")

st.title("📘 Question Bank Generator from TXT/PDF/Image")

# Step 1: Upload files
uploaded_files = st.file_uploader(
    "Upload TXT, PDF or Image files (PNG/JPG/JPEG)", 
    accept_multiple_files=True, 
    type=["pdf", "txt", "png", "jpg", "jpeg"],
    help="Upload any combination of PDF, TXT, or images."
)

st.markdown("---")

# Section for Image to PDF Conversion
st.header("🖼️ Image to PDF Conversion")
image_files = [f for f in uploaded_files if f.type.startswith("image/")]

if image_files:
    st.info(f"{len(image_files)} image file(s) detected for conversion.")
    if st.button("Convert Uploaded Images to PDF"):
        converted_pdfs = []
        for img_file in image_files:
            try:
                img = Image.open(img_file).convert("RGB")
                pdf_bytes = BytesIO()
                img.save(pdf_bytes, format="PDF")
                pdf_bytes.seek(0)
                converted_pdfs.append((img_file.name + ".pdf", pdf_bytes))
            except Exception as e:
                st.error(f"Failed to convert {img_file.name} to PDF: {e}")
        if converted_pdfs:
            st.session_state["converted_image_pdfs"] = converted_pdfs
            st.success(f"Converted {len(converted_pdfs)} image(s) to PDF!")
            for name, pdf in converted_pdfs:
                st.download_button(f"Download {name}", pdf, file_name=name)

else:
    st.info("No images uploaded for conversion.")

st.markdown("---")

# Section for PDF OCR to Extractable PDF
st.header("📄 PDF OCR (Make PDFs Searchable)")
pdf_files = [f for f in uploaded_files if f.type == "application/pdf"]

if pdf_files:
    st.info(f"{len(pdf_files)} PDF file(s) detected for OCR.")
    # Language selection dropdown
    ocr_lang = st.selectbox("Select OCR languages (plus '+'):", ["eng", "tam+hin+eng", "hin+eng", "tam+eng"], index=1)
    
    if st.button("Convert Uploaded PDFs to Searchable PDFs via OCR"):
        ocr_pdfs = []
        failed_files = []
        for pdf_file in pdf_files:
            try:
                pdf_file.seek(0)
                searchable_pdf = ocr_pdf_to_searchable_pdf(pdf_file.read(), languages=ocr_lang)
                ocr_pdfs.append((pdf_file.name.replace(".pdf", "_searchable.pdf"), searchable_pdf))
            except Exception as e:
                failed_files.append((pdf_file.name, str(e)))
        if ocr_pdfs:
            st.session_state["ocr_pdfs"] = ocr_pdfs
            st.success(f"OCR completed for {len(ocr_pdfs)} PDF(s)!")
            for name, pdf in ocr_pdfs:
                st.download_button(f"Download {name}", pdf, file_name=name)
        if failed_files:
            for fname, err in failed_files:
                st.error(f"OCR failed for {fname}: {err}")
else:
    st.info("No PDFs uploaded for OCR.")

st.markdown("---")

# Section for TXT to PDF Conversion (optional, in case user wants)
txt_files = [f for f in uploaded_files if f.type == "text/plain"]
if txt_files:
    st.header("📝 TXT to PDF Conversion")
    if st.button("Convert Uploaded TXT files to PDF"):
        txt_pdfs = []
        failed_txts = []
        for txt_file in txt_files:
            try:
                txt_file.seek(0)
                txt_content = txt_file.read().decode("utf-8")
                pdf_bytes = BytesIO(convert_txt_to_pdf(txt_content))
                txt_pdfs.append((txt_file.name.replace(".txt", ".pdf"), pdf_bytes))
            except Exception as e:
                failed_txts.append((txt_file.name, str(e)))
        if txt_pdfs:
            st.session_state["txt_pdfs"] = txt_pdfs
            st.success(f"Converted {len(txt_pdfs)} TXT(s) to PDF!")
            for name, pdf in txt_pdfs:
                st.download_button(f"Download {name}", pdf, file_name=name)
        if failed_txts:
            for fname, err in failed_txts:
                st.error(f"TXT to PDF conversion failed for {fname}: {err}")
else:
    st.info("No TXT files uploaded.")

st.markdown("---")

# Step 4: Generate Question Bank
st.header("🧠 Generate Question Bank from Uploaded/Converted PDFs")

# Collect all PDFs to send for question generation
def gather_all_pdfs():
    all_pdfs = []
    # Uploaded PDFs (non-OCRed)
    for pdf_file in pdf_files:
        pdf_file.seek(0)
        all_pdfs.append((pdf_file.name, pdf_file.read(), "application/pdf"))
    # OCRed PDFs
    if "ocr_pdfs" in st.session_state:
        for name, pdf_bytesio in st.session_state["ocr_pdfs"]:
            pdf_bytesio.seek(0)
            all_pdfs.append((name, pdf_bytesio.read(), "application/pdf"))
    # Converted image PDFs
    if "converted_image_pdfs" in st.session_state:
        for name, pdf_bytesio in st.session_state["converted_image_pdfs"]:
            pdf_bytesio.seek(0)
            all_pdfs.append((name, pdf_bytesio.read(), "application/pdf"))
    # Converted TXT PDFs
    if "txt_pdfs" in st.session_state:
        for name, pdf_bytesio in st.session_state["txt_pdfs"]:
            pdf_bytesio.seek(0)
            all_pdfs.append((name, pdf_bytesio.read(), "application/pdf"))
    return all_pdfs

if st.button("Generate Question Bank"):
    all_pdfs = gather_all_pdfs()
    if not all_pdfs:
        st.warning("No PDF files available for question generation. Upload or convert some first.")
    else:
        with st.spinner("Generating question bank, please wait..."):
            try:
                files_payload = [("files", (name, content, mimetype)) for name, content, mimetype in all_pdfs]
                response = requests.post(BACKEND_URL, files=files_payload)
                response.raise_for_status()
                data = response.json()
                content = data.get("content", "")
                questions = data.get("questions", {})

                if "error" in questions:
                    st.error(f"Error from backend: {questions['error']}")
                    st.text_area("Raw Backend Response", questions.get("raw_response", ""), height=300)
                else:
                    st.subheader("📄 Combined Content Extracted (Preview)")
                    st.text_area("", content, height=300)

                    st.subheader("❓ Generated Questions")
                    for section, items in questions.items():
                        st.markdown(f"### {section.replace('_', ' ').title()}")
                        for idx, item in enumerate(items):
                            if isinstance(item, dict):
                                st.write(f"Q{idx+1}: {item.get('question')}")
                                opts = item.get("options", [])
                                for opt in opts:
                                    st.write(f"- {opt}")
                                st.write(f"Answer: {item.get('answer')}")
                            else:
                                st.write(f"Q{idx+1}: {item}")

                    # Offer downloads
                    docx_buffer = download_docx(questions)
                    st.download_button("Download Questions as DOCX", docx_buffer, file_name="questions.docx")
                    txt_buffer = download_txt(questions)
                    st.download_button("Download Questions as TXT", txt_buffer, file_name="questions.txt")

            except Exception as e:
                st.error(f"Failed to generate questions: {e}")

st.markdown("---")

# Step 5: Auto Merge All PDFs and Download
st.header("📚 Download All PDFs (Uploaded + Converted + OCRed) Merged")

if st.button("Merge All PDFs and Download"):
    all_pdfs = gather_all_pdfs()
    if not all_pdfs:
        st.warning("No PDFs to merge. Upload or convert PDFs first.")
    else:
        merger = PdfMerger()
        try:
            for _, content, _ in all_pdfs:
                merger.append(BytesIO(content))
            merged_pdf_io = BytesIO()
            merger.write(merged_pdf_io)
            merger.close()
            merged_pdf_io.seek(0)
            st.download_button("Download Merged PDF", merged_pdf_io, file_name="all_combined.pdf", mime="application/pdf")
            st.success("Merged PDF ready for download!")
        except Exception as e:
            st.error(f"Failed to merge PDFs: {e}")
else:
    st.info("Click the button above to merge all PDFs and download.")

